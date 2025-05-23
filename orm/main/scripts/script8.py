import os
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import customtkinter as ctk
from tkinter import filedialog
import time
from json import JSONDecodeError, loads, dumps
from main.models import Empresa, Cuenta, Agencia
from openai import OpenAI


def deserializar_referencias(cuenta: Cuenta) -> list:

    try:
        return loads(cuenta.referencias or "[]")
    except JSONDecodeError:
        return []

def serializar_referencias(referencias: list) -> str:

    return dumps(referencias, ensure_ascii=False)

def searchEmpresa(ppto):
    ppto = ppto.upper()
    pptoFilter = ppto.split()[0]
    empresa = Empresa.objects.filter(key = pptoFilter).first()
    return empresa

def searchAccountByName(addressee, cuentas):
    account = None
    namesClient = cuentas.values_list('nombre', flat=True)
    for name in namesClient:
        if name in addressee:
            account = cuentas.filter(nombre = name).first()
            return account
    return account

def searchAccountByKey(addresseeFilter, cuentas):
    account = None
    keysClient = cuentas.values_list('key', flat=True)
    for key in keysClient:
        if key in addresseeFilter:
            account = cuentas.get(key = key)
            return account
    return account

def searchAccountByReferences(addresseeFilter, cuentas):
    account = None
    for cuenta in cuentas:
        referencias = deserializar_referencias(cuenta)
        for referencia in referencias:
            if referencia in addresseeFilter:
                account = cuentas.get(nombre = cuenta.nombre)
                return account
    return account

def searchCuenta(addressee, client):
    addressee = addressee.upper()
    cuentas = Cuenta.objects.filter(empresa = client)
    cuenta = None
    cuenta = searchAccountByName(addressee, cuentas)
    if cuenta is None:
        addresseeFilter = addressee.split()
        cuenta = searchAccountByKey(addresseeFilter, cuentas)
        if cuenta is None:
            cuenta = searchAccountByReferences(addresseeFilter, cuentas)
    return cuenta

def searchAgencia(addressee, cuenta):
    addressee = addressee.upper()
    agencias = Agencia.objects.filter(cuenta = cuenta)
    agencia = None
    agencia = searchAccountByName(addressee, agencias)
    if agencia is None:
        addresseeFilter = addressee.split()
        agencia = searchAccountByKey(addresseeFilter, agencias)
        if agencia is None:
            agencia = searchAccountByReferences(addresseeFilter, agencias)
    return agencia

def obtenerCuenta(destinatarios, nombre_cotizacion, empresa):
    cuenta = searchCuenta(destinatarios, empresa)
    if cuenta is None:
        cuenta = searchCuenta(nombre_cotizacion, empresa)
    return cuenta

def obtenerAgencia(destinatarios, nombre_cotizacion, cuenta):
    agencia = searchAgencia(destinatarios, cuenta)
    if agencia is None:
        agencia = searchAgencia(nombre_cotizacion, cuenta)
    return agencia


def sortImagesByDate(folderImages, invert=False):
    images = sorted(
        [f for f in os.listdir(folderImages) if f.lower().endswith(('png', 'jpg', 'jpeg', 'gif', 'bmp'))],
        key=lambda x: os.path.getmtime(os.path.join(folderImages, x)),
        reverse=invert
    )
    return images

def sortImagesByName(folderImages, invert=False):
    images = sorted(
        [f for f in os.listdir(folderImages) if f.lower().endswith(('png', 'jpg', 'jpeg', 'gif', 'bmp'))],
        reverse=invert
    )
    return images

def getImages(folderImages, sortByDate=True, invert=False):
    if sortByDate:
        return sortImagesByDate(folderImages, invert)
    else:
        return sortImagesByName(folderImages, invert)

def createTable(table, images, folderImages, heightRowImages=6, widthRow=7.8, heightRowDescripcion=0.28, heightImage=5.5):
    table.autofit = False
    
    for i, image in enumerate(images): 
        if i % 2 == 0:
            rowImages = table.add_row()
            rowImages.height = Cm(heightRowImages)
            for cell in rowImages.cells:
                cell.width = Cm(widthRow)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        pathImage = os.path.join(folderImages, image)
        cell = rowImages.cells[i % 2]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(pathImage, height=Cm(heightImage))
        
        if i % 2 == 1 or i == len(images) - 1:
            rowDescription = table.add_row()
            rowDescription.height = Cm(heightRowDescripcion)
            for cell in rowDescription.cells:
                cell.width = Cm(widthRow)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.text = ""
                

def insertTable(doc, images, folderImages, marker = "{{TABLAIMAGENES}}"):
    for i, paragraph in enumerate(doc.paragraphs):
        if marker in paragraph.text:
            p = doc.paragraphs[i]
            p.text = p.text.replace(marker, '')
            table = doc.add_table(rows=0, cols=2, style='Table Grid')
            createTable(table, images, folderImages)
            tbl = p._p.addnext(table._tbl)
            break

def insertDate(doc):
    hoy = date.today().strftime("%d-%m-%Y")

    def reemplazar_en_runs(runs):
        for run in runs:
            if "{{FECHA}}" in run.text:
                run.text = run.text.replace("{{FECHA}}", hoy)

    for p in doc.paragraphs:
        reemplazar_en_runs(p.runs)   
        
def insertAgencia(doc, agencia):

    def reemplazar_en_runs(runs):
        for run in runs:
            if "{{AGENCIA}}" in run.text:
                run.text = run.text.replace("{{AGENCIA}}", agencia)

    for p in doc.paragraphs:
        reemplazar_en_runs(p.runs)   
        
from openpyxl import load_workbook
import re
        
def leerDatosExcel(ruta):
    wb = load_workbook(ruta, data_only=True, read_only=True)
    sheet = wb.active
    
    contenido = []
    
    for row in sheet.iter_rows(values_only=True):
        if any(cell is not None and str(cell).strip() != "" for cell in row):
            fila = [str(cell).strip() if cell is not None else "" for cell in row]
            contenido.append("\t".join(fila))
    
    contenido= "\n".join(contenido)

    datos = {
        "agenciaDestinatarios": sheet['D17'].value,
        "referencia": sheet['D19'].value,
        "ppto": sheet['J13'].value,
        "contenido": contenido,
    }
    return datos

def leerDatosTottus(ruta, nameCotizacion):
    wb = load_workbook(ruta, data_only=True, read_only=True)
    sheet = wb.active
    
    contenido = []
    
    for row in sheet.iter_rows(values_only=True):
        if any(cell is not None and str(cell).strip() != "" for cell in row):
            fila = [str(cell).strip() if cell is not None else "" for cell in row]
            contenido.append("\t".join(fila))
    
    contenido= "\n".join(contenido)
    
    match = re.search(r'TT 00\d{2}', nameCotizacion)
    ppto = match.group(0) if match else nameCotizacion.split(".")[0]
    
    datos = {
        "agenciaDestinatarios": nameCotizacion,
        "referencia": sheet['C10'].value,
        "ppto": ppto,
        "contenido": contenido
    }
    return datos

    
def extractData(cotizacion, nameCotizacion):
    datos_excel = leerDatosExcel(cotizacion)
    if not datos_excel['agenciaDestinatarios']:
        datos_excel = leerDatosTottus(cotizacion, nameCotizacion)
        
    empresa = searchEmpresa(datos_excel['ppto'])
    cuenta = obtenerCuenta(datos_excel['agenciaDestinatarios'], nameCotizacion, empresa)
    agencia = obtenerAgencia(datos_excel['agenciaDestinatarios'], nameCotizacion, cuenta)


    return agencia, datos_excel['referencia'], datos_excel["contenido"]
    
def mainIBK(folderPhotos, agencia, ordenarPorFecha = True, invertir = False, datos =""):
    archivo_entrada = "FORMATO INFORME GUSTAVO.docx"
    archivo_salida = "formato_con_tabla.docx"

    if not os.path.exists(archivo_entrada):
        print(f"❌ No se encontró el archivo Word: {archivo_entrada}")
        return

    if not os.path.exists(folderPhotos):
        print(f"❌ No se encontró la carpeta de imágenes: {folderPhotos}")
        return

    doc = Document(archivo_entrada)
    imagenes = getImages(folderPhotos, sortByDate=ordenarPorFecha, invert=invertir)

    if not imagenes:
        print("⚠️ No se encontraron imágenes válidas en la carpeta.")
        return
    
    insertTable(doc, imagenes, folderPhotos)
    insertDate(doc=doc)  
    insertAgencia(doc, agencia)

    doc.save(archivo_salida)
    print(f"✅ Documento generado: {archivo_salida}")
    

def obtener_carpeta_fotos(ruta_base):
    carpeta_fotos = os.path.join(ruta_base, "FOTOS")
    if os.path.isdir(carpeta_fotos):
        return carpeta_fotos
    return None

def obtenerCotizacion(ruta_folder):
    path_cotizacion_excel = None
    nombre_cotizacion_excel = None

    if not os.path.exists(ruta_folder):
        print("❌ La ruta no existe")
        return None, None

    for item in os.listdir(ruta_folder):
        path = os.path.join(ruta_folder, item)
        if not os.path.isfile(path):
            continue
        nombre = item.lower()

        if (nombre.endswith(".xlsx") and 
            "cotiza" in nombre):
            path_cotizacion_excel = path
            nombre_cotizacion_excel = item
            break

    return path_cotizacion_excel, nombre_cotizacion_excel

def getAlcances(datos):
    apiKey = "sk-918a3a5bbc1b4d4e97032576083b2d04"
    client = OpenAI(api_key=apiKey, base_url="https://api.deepseek.com")
    response = client.chat.completions.create(
    model="deepseek-chat",
    messages=[
        {
            "role": "system",
            "content": "Eres un asistente que genera un título de servicio y alcances técnicos para informes de mantenimiento. Devuelve ÚNICAMENTE un array JSON de 2 strings: [servicio, alcances]. Servicio debe ser un título conciso (ej: 'mantenimiento de [equipo] Nro. [X]'). Alcances debe ser una lista de actividades en infinitivo separadas por comas, sin viñetas."
        },
        {
            "role": "user",
            "content": f"Extrae título de servicio y alcances de: '{datos}'"
        }
    ],
    stream=False
)
    return response.choices[0].message.content
    
def procesarCarpeta(ruta, obervaciones):
    fotos = obtener_carpeta_fotos(ruta)
    cotizacion, nameCotizacion = obtenerCotizacion(ruta)
    agencia, referencia, datos = extractData(cotizacion=cotizacion, nameCotizacion=nameCotizacion)
    
    mainIBK(folderPhotos=fotos, agencia=agencia.nombre, datos=datos)
    


def run_gui():
    def seleccionar_carpeta():
        ruta = filedialog.askdirectory()
        if ruta:
            folder_var.set(ruta)

    def ejecutar_proceso():
        ruta = folder_var.get()
        observaciones = observaciones_entry.get().strip()

        if not ruta or not os.path.exists(ruta):
            resultado_var.set("⚠️ Carpeta no válida.")
            return

        inicio = time.perf_counter()
        procesarCarpeta(ruta, observaciones)
        fin = time.perf_counter()
        duracion = fin - inicio
        resultado_var.set(f"✅ Listo en {duracion:.2f} segundos.")

    ctk.set_appearance_mode("system")
    ctk.set_default_color_theme("green")

    app = ctk.CTk()
    app.title("Generar Informe")
    app.geometry("600x250")

    folder_frame = ctk.CTkFrame(app)
    folder_frame.pack(pady=10, padx=20, fill="x")

    folder_var = ctk.StringVar()
    ctk.CTkLabel(folder_frame, text="Carpeta:").pack(side="left", padx=5)
    ctk.CTkEntry(folder_frame, textvariable=folder_var, width=400).pack(side="left", padx=5)
    ctk.CTkButton(folder_frame, text="Seleccionar", command=seleccionar_carpeta).pack(side="left")

    ctk.CTkLabel(app, text="Observaciones:").pack(pady=(10, 0))
    observaciones_entry = ctk.CTkEntry(app, width=550)
    observaciones_entry.pack(pady=5)

    ctk.CTkButton(app, text="Procesar y Generar", command=ejecutar_proceso).pack(pady=10)

    resultado_var = ctk.StringVar()
    resultado_label = ctk.CTkLabel(app, textvariable=resultado_var, text_color="gray")
    resultado_label.pack(pady=5)

    app.mainloop()

def run():
    run_gui()